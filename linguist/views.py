from django.shortcuts import render, HttpResponse  
from django.http import JsonResponse , FileResponse
from deep_translator import GoogleTranslator
from django.core.files.base import ContentFile
import json
from django.views.decorators.csrf import csrf_protect
import PyPDF2
import os
from docx import Document
import speech_recognition as sr
from pydub import AudioSegment 
from docx2pdf import convert



@csrf_protect

def index(req):
    languages = GoogleTranslator().get_supported_languages(as_dict=True)
    languages = [ (code , language.title()) for code , language  in languages.items() ]

    data = {'message': languages}  
    return render(req , 'index.html' , data)

      

def tran(req):
     if req.method == 'POST':
           
        data = json.loads(req.body)

        txt = data['txt']
        lang = data['lang'].lower()

        translation = GoogleTranslator(source='auto', target=lang).translate(txt)
        return JsonResponse({'message': translation})


def doc(req):
    def pdfReader(pdfFile , translator):
        
        pdf_reader = PyPDF2.PdfReader(pdfFile)
        text = ""

        for page in pdf_reader.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + "\n"
    
        max_chunk_size = 500
        chunks = [text[i:i + max_chunk_size] for i in range(0, len(text), max_chunk_size)]
        translated_chunks = []
        
        for chunk in chunks:           
            translated = translator.translate(chunk)
            translated_chunks.append(translated)

        translatedTxt = "".join(translated_chunks)

        return translatedTxt
    
    def txtReader(txtfile, translator):
       
        textR = txtfile.read()
        text = textR.decode('utf-8')
        

        max_chunk_size = 4999
        chunks = [text[i:i + max_chunk_size] for i in range(0, len(text), max_chunk_size)]
        translated_chunks = [translator.translate(chunk) for chunk in chunks]

        content = "".join(translated_chunks)

        try:
            fileName = "sample.txt"
            path = rf"temp/{fileName}"
            with open(path, 'w' , encoding='utf-8') as file:
                file.writelines(content)
            os.remove(path)
        except Exception as e:
            print(f"An error occurred: {e}")
        
        # return FileResponse(path , as_attachment=True , filename=fileName)


        return content
        

    def docxReader(docFile , translator):
        try:
            doc = Document(docFile)
            text = "".join([paragraph.text for paragraph in doc.paragraphs])
           
          
            max_chunk_size = 4999
            chunks = [text[i:i + max_chunk_size] for i in range(0, len(text), max_chunk_size)]
            translated_chunks = [translator.translate(chunk) for chunk in chunks]

            
            return "\n".join(translated_chunks)
        
        except Exception as e:
            return f"Error reading DOCX do: {str(e)}"
    

    def docxWriter(newTxt , f):

        newDoc = Document()
        newDoc.add_paragraph(newTxt)
        fileName = ''
        path = ''
        if f == 0:
            fileName = "sample.pdf"
            path = rf"temp/{fileName}"
            os.remove(path)
            newDoc.save(r"temp/pdfSample.docx")
            convert(r"temp/pdfSample.docx",path)

            return FileResponse(open(path , 'rb') , as_attachment=True , filename=fileName)

        else:
            fileName = "sample.docx"
            path = rf"temp/{fileName}"
            os.remove(path)
            newDoc.save(path)

            return FileResponse(path , as_attachment=True , filename=fileName)

    def audioTxt(mp3File , translator):
           
        mp3_file = mp3File
        wav_file = r"temp/audio.wav"
        os.remove(wav_file)

        sound = AudioSegment.from_file(mp3_file, format="mp3")
        sound = sound.set_frame_rate(16000).set_channels(1)  

        sound = sound + 5  

        sound.export(wav_file, format="wav")

        r = sr.Recognizer()
        with sr.AudioFile(wav_file) as source:
            audio_data = r.record(source)

        try:
            text = r.recognize_google(audio_data)
            translatedTxt = translator.translate(text)
            return translatedTxt
        except sr.UnknownValueError:
            return "Speech Recognition could not understand audio"
        except sr.RequestError as e:
            return f"Could not request results from Speech Recognition service; {e}"

        

    lang = req.POST['lang'].lower()
    file = req.FILES['file']

    file_name = file.name
    file_extension = os.path.splitext(file_name)[1].lower()
    
    translator = GoogleTranslator(source='auto', target=lang)

    f = ''

    match file_extension:
        case '.pdf':
            f = 0
            txt = pdfReader(file , translator)
            docxWriter(txt , f)
        case '.txt':      
            txtReader(file , translator )
        case '.docx':
            f = 2
            txt = docxReader(file , translator)
            docxWriter(txt , f)
        case '.mp3':
            txt = audioTxt(file , translator) 
            data = { 
            # "txt" : text,
            "message": txt
            }
            return JsonResponse(data)

        case default:
            txt = "Not Valid Document , Please Upload Valid Document ( '.docx', '.pdf', '.txt', '.mp3' )"


     
    

